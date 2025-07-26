import streamlit as st
from docx import Document
from datetime import datetime
import io
import smtplib
from email.message import EmailMessage

# Função para gerar o contrato preenchido
def gerar_contrato(dados):
    doc = Document("Contrato_Modelo_Com_Placeholders_Novo.docx")

    for p in doc.paragraphs:
        for chave, valor in dados.items():
            if chave in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if chave in inline[i].text:
                        inline[i].text = inline[i].text.replace(chave, valor)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for chave, valor in dados.items():
                    if chave in cell.text:
                        cell.text = cell.text.replace(chave, valor)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# Função para enviar o contrato por e-mail
def enviar_email(arquivo, nome_cliente):
    email_de = "contratosguilherme.enviador@gmail.com"
    senha = "quprkjbbttuyxwnv"
    email_para = "contratosguilherme.enviador@gmail.com"

    msg = EmailMessage()
    msg["Subject"] = f"Contrato de Honorários - {nome_cliente}"
    msg["From"] = email_de
    msg["To"] = email_para
    msg.set_content(f"Segue em anexo o contrato gerado para {nome_cliente}.")

    msg.add_attachment(arquivo.getvalue(),
                       maintype='application',
                       subtype='vnd.openxmlformats-officedocument.wordprocessingml.document',
                       filename="Contrato_Gerado.docx")

    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
        smtp.login(email_de, senha)
        smtp.send_message(msg)

st.title("Gerador de Contrato de Honorários")

st.markdown("Preencha as informações abaixo para gerar automaticamente o contrato personalizado em Word.")

# Formulário de entrada
dados = {}
dados["{{CONTRATANTE_NOME}}"] = st.text_input("Nome do Contratante", "João da Silva")
dados["{{CPF}}"] = st.text_input("CPF", "123.456.789-00")
dados["{{RG}}"] = st.text_input("RG", "MG-1234567")
dados["{{EMAIL}}"] = st.text_input("Email", "joao@email.com")
dados["{{ENDERECO}}"] = st.text_area("Endereço completo", "Rua Exemplo, nº 100, Centro, São Paulo - SP")
dados["{{OBJETO}}"] = st.text_area("Objeto do Contrato", "Ação de cobrança nº 1036865-65.2024.8.26.0001")
dados["{{DATA_ASSINATURA}}"] = st.date_input("Data de Assinatura", datetime.today()).strftime("%d de %B de %Y")

# Tabela de honorários
st.markdown("### Parcelas de Honorários")
parcelas = []
num_parcelas = st.number_input("Número de parcelas", min_value=1, max_value=36, value=3)

for i in range(num_parcelas):
    valor = st.number_input(f"Valor da {i+1}ª parcela", min_value=0.0, value=1000.0, key=f"valor_{i}")
    venc = st.date_input(f"Vencimento da {i+1}ª parcela", key=f"venc_{i}")
    parcelas.append(f"{i+1}ª Parcela\tR${valor:.2f}\t{venc.strftime('%d/%m/%Y')}")

dados["{{TABELA_PARCELAS}}"] = "\n".join(parcelas)

# Geração do contrato
if st.button("Gerar Contrato"):
    contrato = gerar_contrato(dados)
    enviar_email(contrato, dados["{{CONTRATANTE_NOME}}"])
    st.success("Contrato gerado e enviado por e-mail com sucesso!")
    st.download_button("Clique para baixar o contrato Word", data=contrato, file_name="Contrato_Gerado.docx")
